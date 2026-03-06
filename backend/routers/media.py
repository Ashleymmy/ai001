"""Media routes: /api/upload-reference, /api/upload, /api/uploads/*, /api/images/*, /api/videos/ref/*, /api/studio/parse-document."""

import os
import base64
from typing import Optional, List
from fastapi import APIRouter, HTTPException, UploadFile, File, Request
from fastapi.responses import FileResponse
from pydantic import BaseModel

from services.storage_service import storage
import dependencies as deps

router = APIRouter(prefix="/api", tags=["media"])


class DeleteImagesRequest(BaseModel):
    ids: List[str]


@router.post("/upload-reference")
async def upload_reference(file: UploadFile = File(...)):
    if not file.content_type or not file.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="只支持图片文件")
    content = await file.read()
    base64_data = base64.b64encode(content).decode("utf-8")
    mime_type = file.content_type
    return {"dataUrl": f"data:{mime_type};base64,{base64_data}", "filename": file.filename}


@router.post("/upload")
async def upload_file(request: Request, file: UploadFile = File(...)):
    import time
    content_type = file.content_type or 'application/octet-stream'
    ext = os.path.splitext(file.filename or '')[1].lower()
    category = 'unknown'
    max_size = 10 * 1024 * 1024
    if content_type in deps.ALLOWED_FILE_TYPES:
        file_config = deps.ALLOWED_FILE_TYPES[content_type]
        category = file_config['category']
        max_size = file_config['max_size']
    elif content_type.startswith('image/'):
        category = 'image'
        max_size = 20 * 1024 * 1024
    elif content_type.startswith('video/'):
        category = 'video'
        max_size = 100 * 1024 * 1024
    elif content_type.startswith('audio/'):
        category = 'audio'
        max_size = 25 * 1024 * 1024
    elif content_type.startswith('text/') or ext in ['.py', '.js', '.ts', '.jsx', '.tsx', '.java', '.cpp', '.c', '.go', '.rs', '.sql', '.yaml', '.yml']:
        category = 'code'
        max_size = 10 * 1024 * 1024

    content = await file.read()
    file_size = len(content)
    if file_size > max_size:
        raise HTTPException(status_code=400, detail=f"文件过大，最大允许 {max_size // 1024 // 1024}MB")

    timestamp = int(time.time() * 1000)
    safe_filename = f"{timestamp}_{file.filename}"
    category_dir = os.path.join(deps.UPLOAD_DIR, category)
    os.makedirs(category_dir, exist_ok=True)
    file_path = os.path.join(category_dir, safe_filename)
    with open(file_path, 'wb') as f:
        f.write(content)

    file_url = f"/api/uploads/{category}/{safe_filename}"
    absolute_url = f"{str(request.base_url).rstrip('/')}{file_url}"

    preview_url = None
    if category == 'image':
        base64_data = base64.b64encode(content).decode("utf-8")
        preview_url = f"data:{content_type};base64,{base64_data}"

    text_content = None
    if category in ['code', 'document'] and file_size < 1024 * 1024 and ext not in ['.pdf', '.docx']:
        try:
            import codecs
            if content.startswith(codecs.BOM_UTF8):
                text_content = content.decode('utf-8-sig')
            elif content.startswith(codecs.BOM_UTF16_LE) or content.startswith(codecs.BOM_UTF16_BE):
                text_content = content.decode('utf-16')
            else:
                try:
                    text_content = content.decode('utf-8')
                except UnicodeDecodeError:
                    sample = content[:4096]
                    if sample and sample.count(b'\x00') / len(sample) > 0.2:
                        try:
                            text_content = content.decode('utf-16-le')
                        except UnicodeDecodeError:
                            try:
                                text_content = content.decode('utf-16-be')
                            except UnicodeDecodeError:
                                text_content = None
                    if text_content is None:
                        try:
                            text_content = content.decode('gb18030')
                        except UnicodeDecodeError:
                            try:
                                text_content = content.decode('gbk')
                            except UnicodeDecodeError:
                                text_content = content.decode('utf-8', errors='replace')
        except Exception:
            pass

    if ext == '.docx' and text_content is None:
        try:
            from docx import Document
            from io import BytesIO
            doc = Document(BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text_content = '\n'.join(paragraphs)
        except Exception as e:
            print(f"[Upload] 解析 Word 文档失败: {e}")

    if ext == '.pdf' and text_content is None:
        try:
            from PyPDF2 import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            pages_text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)
            text_content = '\n\n'.join(pages_text)
        except Exception as e:
            print(f"[Upload] 解析 PDF 文档失败: {e}")

    print(f"[Upload] 文件已上传: {file.filename} -> {file_path} ({file_size} bytes, {category})")

    return {
        "success": True,
        "file": {
            "id": f"upload_{timestamp}",
            "name": file.filename,
            "size": file_size,
            "type": content_type,
            "category": category,
            "url": file_url,
            "absoluteUrl": absolute_url,
            "previewUrl": preview_url,
            "content": text_content
        }
    }


@router.post("/studio/parse-document")
async def studio_parse_document(file: UploadFile = File(...)):
    content = await file.read()
    ext = os.path.splitext(file.filename or '')[1].lower()
    text_content = None

    if ext in ['.txt', '.md', '.markdown']:
        try:
            import codecs
            if content.startswith(codecs.BOM_UTF8):
                text_content = content.decode('utf-8-sig')
            elif content.startswith(codecs.BOM_UTF16_LE) or content.startswith(codecs.BOM_UTF16_BE):
                text_content = content.decode('utf-16')
            else:
                try:
                    text_content = content.decode('utf-8')
                except UnicodeDecodeError:
                    try:
                        text_content = content.decode('gb18030')
                    except UnicodeDecodeError:
                        try:
                            text_content = content.decode('gbk')
                        except UnicodeDecodeError:
                            text_content = content.decode('utf-8', errors='replace')
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"文本文件解码失败: {e}")
    elif ext == '.docx':
        try:
            from docx import Document as DocxDocument
            from io import BytesIO
            doc = DocxDocument(BytesIO(content))
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            text_content = '\n'.join(paragraphs)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Word 文档解析失败: {e}")
    elif ext == '.pdf':
        try:
            from PyPDF2 import PdfReader
            from io import BytesIO
            reader = PdfReader(BytesIO(content))
            pages_text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)
            text_content = '\n\n'.join(pages_text)
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"PDF 文档解析失败: {e}")
    else:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式: {ext}，支持 .txt/.md/.docx/.pdf")

    if text_content is None:
        raise HTTPException(status_code=400, detail="文档内容为空或无法提取")
    print(f"[ParseDocument] 解析文档: {file.filename} ({ext}), 提取 {len(text_content)} 字符")
    return {"text": text_content, "filename": file.filename}


@router.get("/uploads/{category}/{filename}")
async def get_uploaded_file(category: str, filename: str):
    file_path = os.path.join(deps.UPLOAD_DIR, category, filename)
    if not os.path.exists(file_path):
        raise HTTPException(status_code=404, detail="文件不存在")
    ext = os.path.splitext(filename)[1].lower()
    mime_types = {
        '.jpg': 'image/jpeg', '.jpeg': 'image/jpeg', '.png': 'image/png',
        '.gif': 'image/gif', '.webp': 'image/webp',
        '.mp4': 'video/mp4', '.webm': 'video/webm', '.mov': 'video/quicktime',
        '.mp3': 'audio/mpeg', '.wav': 'audio/wav', '.m4a': 'audio/mp4',
        '.pdf': 'application/pdf', '.txt': 'text/plain', '.md': 'text/markdown',
        '.json': 'application/json', '.xml': 'application/xml',
        '.html': 'text/html', '.css': 'text/css', '.js': 'text/javascript',
    }
    media_type = mime_types.get(ext, 'application/octet-stream')
    return FileResponse(file_path, media_type=media_type)


@router.get("/images/history")
async def get_image_history(limit: int = 100, project_id: Optional[str] = None):
    images = storage.list_generated_images(limit, project_id=project_id)
    return {"images": images}


@router.get("/agent/images/history")
async def get_agent_image_history(limit: int = 100, project_id: Optional[str] = None):
    images = storage.list_agent_generated_images(limit, project_id=project_id)
    return {"images": images}


@router.delete("/images/history/{image_id}")
async def delete_image_history(image_id: str):
    success = storage.delete_generated_image(image_id)
    if not success:
        raise HTTPException(status_code=404, detail="图像记录不存在")
    return {"status": "ok"}


@router.post("/images/history/delete-batch")
async def delete_images_batch(request: DeleteImagesRequest):
    deleted = storage.delete_generated_images_batch(request.ids)
    return {"status": "ok", "deleted": deleted}


@router.get("/images/ref/{filename}")
async def get_reference_image(filename: str):
    filepath = os.path.join(deps.REF_IMAGES_DIR, filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="图片不存在")
    return FileResponse(filepath, media_type="image/png")


@router.get("/videos/ref/{filename}")
async def get_video_reference_image(filename: str):
    from services.video_service import VIDEO_DIR
    filepath = os.path.join(VIDEO_DIR, filename)
    if not os.path.exists(filepath):
        raise HTTPException(status_code=404, detail="图片不存在")
    ext = filename.split(".")[-1].lower()
    media_type = "image/png" if ext == "png" else "image/jpeg"
    return FileResponse(filepath, media_type=media_type)
